"""Core IR / adapter / worker-protocol / native-lane regression checks."""

import conftest_path  # noqa: F401

import base64
import io
import json
import pathlib
import sys
import tempfile
import time
import zipfile

import numpy as np
import cv2

from core import doc_ir, auto_ocr, markup_detect as md

WORK = pathlib.Path(tempfile.mkdtemp(prefix="textlab_test_"))


def crop_b64(kind):
    img = np.full((40, 40, 3), 255, np.uint8)
    cv2.rectangle(img, (2, 2), (37, 37), (0, 0, 0), 1)
    if kind == "checked":
        cv2.line(img, (8, 8), (31, 31), (0, 0, 0), 3)
        cv2.line(img, (31, 8), (8, 31), (0, 0, 0), 3)
    ok, enc = cv2.imencode(".png", img)
    return base64.b64encode(enc.tobytes()).decode()


PAGE_JSON = {
    "page_number": 1, "width": 800, "height": 1000,
    "parsing_res_list": [
        {"block_label": "title", "block_content": "Survey Form", "block_bbox": [10, 10, 400, 40], "block_order": 0},
        {"block_label": "checkbox", "block_content": "", "block_bbox": [10, 90, 50, 130], "block_order": 1},
        {"block_label": "table", "block_content": "<table><tr><th>A</th><th>B</th></tr><tr><td>1</td><td>2</td></tr></table>", "block_bbox": [10, 200, 400, 300], "block_order": 2},
        {"block_label": "formula", "block_content": "E = mc^2", "block_bbox": [10, 320, 200, 360], "block_order": 3},
    ],
    "layout_det_res": [{"label": "checkbox", "score": 0.85, "coordinate": [10, 90, 50, 130]}],
    "assets": {"1": {"b64": crop_b64("checked"), "ext": "png"}},
    "markdown": "# Survey",
}


def test_adapter_and_exports():
    page = doc_ir.from_paddle_vl(PAGE_JSON)
    assert [r.type for r in page.regions] == ["title", "checkbox", "table", "formula"]
    doc = doc_ir.Document(pages=[page], source_name="s.pdf")
    tables = doc_ir.tables_to_dataframes(doc)
    assert len(tables) == 1 and list(tables[0]["dataframe"].columns) == ["A", "B"]
    mdtxt = doc_ir.to_markdown(doc)
    assert "## Survey Form" in mdtxt and "$$" in mdtxt
    names = zipfile.ZipFile(io.BytesIO(doc_ir.build_full_bundle(doc))).namelist()
    assert any(n.startswith("tables/") for n in names)
    assert any(n.startswith("assets/") for n in names)
    json.loads(doc_ir.to_json(doc))


def test_text_export():
    doc = doc_ir.Document(pages=[doc_ir.from_paddle_vl(PAGE_JSON)], source_name="s.pdf")
    txt = doc_ir.to_text(doc)
    assert "Survey Form" in txt and "E = mc^2" in txt
    # Tables become readable columns, never raw markup.
    assert "A" in txt and "1" in txt
    assert "<table>" not in txt and "<td>" not in txt and "##" not in txt
    # A checkbox keeps its state; the crop is named rather than embedded.
    assert "[?]" in txt or "[x]" in txt or "[ ]" in txt

    two = doc_ir.Document(pages=[
        doc_ir.from_paddle_vl(PAGE_JSON), doc_ir.from_paddle_vl(PAGE_JSON)
    ])
    two.pages[1].page_number = 2
    assert "--- page 2 ---" in doc_ir.to_text(two)
    assert doc_ir.to_text(doc_ir.Document(pages=[])).strip() == ""


def test_docx_export():
    doc = doc_ir.Document(pages=[doc_ir.from_paddle_vl(PAGE_JSON)], source_name="s.pdf")
    blob = doc_ir.build_docx(doc, "s")
    if blob is None:  # python-docx absent: the caller hides the download
        return
    assert "word/document.xml" in zipfile.ZipFile(io.BytesIO(blob)).namelist()

    # A figure carries its crop into the document; a checkbox stays a text
    # marker, so embedding is checked on a page that actually has a figure.
    with_figure = doc_ir.Document(pages=[doc_ir.Page(page_number=1, regions=[
        doc_ir.Region("r1", doc_ir.FIGURE, [0, 0, 40, 40], 0, {"text": "Fig 1"},
                      asset={"b64": crop_b64("checked"), "ext": "png"}),
    ])])
    fig_blob = doc_ir.build_docx(with_figure, "f")
    fig_names = zipfile.ZipFile(io.BytesIO(fig_blob)).namelist()
    assert any(n.startswith("word/media/") for n in fig_names), "figure crop not embedded"

    import docx as _docx

    # The caption is italic on its run. Setting ``Paragraph.italic`` instead is
    # accepted silently by python-docx and formats nothing.
    caption_runs = [
        run
        for paragraph in _docx.Document(io.BytesIO(fig_blob)).paragraphs
        for run in paragraph.runs
        if run.text == "Fig 1"
    ]
    assert caption_runs, "figure caption missing"
    assert all(run.italic for run in caption_runs), "figure caption is not italic"

    parsed = _docx.Document(io.BytesIO(blob))
    assert parsed.paragraphs[0].style.name.startswith("Heading")
    assert parsed.paragraphs[0].text == "Survey Form"
    # The table must be a real Word table, not a pasted string.
    assert len(parsed.tables) == 1
    table = parsed.tables[0]
    assert (len(table.rows), len(table.columns)) == (2, 2)
    assert [c.text for c in table.rows[0].cells] == ["A", "B"]
    assert [c.text for c in table.rows[1].cells] == ["1", "2"]
    assert doc_ir.build_docx(doc_ir.Document(pages=[]), "empty") is not None


def test_full_bundle_carries_every_format():
    doc = doc_ir.Document(pages=[doc_ir.from_paddle_vl(PAGE_JSON)], source_name="s.pdf")
    names = zipfile.ZipFile(io.BytesIO(doc_ir.build_full_bundle(doc, "s"))).namelist()
    for expected in ("s.md", "s.txt", "s.json"):
        assert expected in names, f"{expected} missing from bundle: {names}"
    if doc_ir.build_docx(doc, "s") is not None:
        assert "s.docx" in names


def test_batch_outputs_match_the_single_document_downloads():
    """A batch result must carry every format the single-file page offers."""
    doc = doc_ir.Document(pages=[doc_ir.from_paddle_vl(PAGE_JSON)], source_name="s.pdf")
    doc.searchable_pdf = b"%PDF-1.7 fake"
    out = WORK / "batch_out"
    written = doc_ir.write_document_outputs(doc, out, "document")

    bundle = zipfile.ZipFile(io.BytesIO(doc_ir.build_full_bundle(doc, "document")))
    def kinds(names):
        return {
            pathlib.PurePath(n).suffix or pathlib.PurePath(n).name
            for n in names if not n.endswith("/")
        }
    missing = kinds(bundle.namelist()) - kinds(written)
    assert not missing, f"batch is missing formats the bundle has: {missing}"

    for expected in ("document.md", "document.txt", "document.json",
                     "document_searchable.pdf", "models_used.txt"):
        assert (out / expected).exists(), f"{expected} not written: {written}"

    # Batch writes one merged summary at the root instead of repeating it in
    # every per-file folder.
    per_file = doc_ir.write_document_outputs(
        doc, WORK / "batch_no_prov", "document", provenance=False)
    assert "models_used.txt" not in per_file, per_file
    assert not (WORK / "batch_no_prov" / "models_used.txt").exists()
    assert any(n.startswith("tables/") for n in written), written
    assert any(n.startswith("assets/") for n in written), written
    if doc_ir.build_docx(doc, "document") is not None:
        assert (out / "document.docx").exists()
    # Written where they are claimed, and not empty.
    assert (out / "document.txt").read_text(encoding="utf-8").strip()
    assert (out / "document_searchable.pdf").read_bytes() == b"%PDF-1.7 fake"


def test_model_provenance_is_citable_and_derived():
    """The citation summary must come from what ran, not a restated constant."""
    page = doc_ir.from_paddle_vl(PAGE_JSON)
    doc = doc_ir.Document(pages=[page], source_name="s.pdf")
    models = doc_ir.model_provenance(doc)
    assert "PaddleOCR-VL 1.6" in " ".join(models["text_recognition"])
    # No figure was described, so no description model may be claimed.
    assert "figure_descriptions" not in models

    figure = doc_ir.Region("f1", doc_ir.FIGURE, [0, 0, 10, 10], 9, {"text": ""},
                           asset={"b64": crop_b64("checked"), "ext": "png"})
    figure.visual_description = doc_ir.VisualDescription(
        description="a chart", source="ollama-local", model="qwen3-vl:30b-a3b-instruct")
    page.regions.append(figure)
    doc.extra_tools["text_layer"] = "Tesseract 4.1.1 (deu), word geometry only"
    models = doc_ir.model_provenance(doc)
    assert models["figure_descriptions"] == ["qwen3-vl:30b-a3b-instruct (ollama-local)"]
    assert "Tesseract 4.1.1" in models["text_layer"]

    # A born-digital page names no recognition model, because none ran.
    native = doc_ir.Document(pages=[doc_ir.Page(page_number=1, source="native")])
    assert "no recognition model" in " ".join(doc_ir.model_provenance(native)["text_recognition"])

    # It travels with the canonical JSON and as readable lines.
    assert "models" in json.loads(doc_ir.to_json(doc))
    text = doc_ir.model_provenance_text(doc)
    assert "Text recognition:" in text and "Figure descriptions:" in text


def test_batch_provenance_is_the_union_over_the_files():
    """A batch is not uniform, so its one summary must cover every file."""
    scanned = {"text_recognition": ["PaddleOCR-VL 1.6"],
               "figure_descriptions": ["qwen3-vl:30b-a3b-instruct (ollama-local)"],
               "text_layer": "Tesseract 4.1.1 (deu), word geometry only"}
    born_digital = {"text_recognition": ["PyMuPDF text extraction (no recognition model)"]}
    another_scan = {"text_recognition": ["PaddleOCR-VL 1.6"],
                    "text_layer": "Tesseract 4.1.1 (eng), word geometry only"}

    merged = doc_ir.merge_provenance([scanned, born_digital, another_scan])
    # Every lane that ran is named, once, in first-seen order.
    assert merged["text_recognition"] == [
        "PaddleOCR-VL 1.6", "PyMuPDF text extraction (no recognition model)"]
    assert merged["figure_descriptions"] == ["qwen3-vl:30b-a3b-instruct (ollama-local)"]
    # A scalar from one file and a different one from another both survive.
    assert merged["text_layer"] == [
        "Tesseract 4.1.1 (deu), word geometry only",
        "Tesseract 4.1.1 (eng), word geometry only"]

    text = doc_ir.provenance_to_text(merged)
    assert "Text recognition: PaddleOCR-VL 1.6, PyMuPDF" in text
    assert doc_ir.merge_provenance([]) == {}
    assert doc_ir.provenance_to_text({}) == ""


def test_finalize_vl_page():
    raster = WORK / "p1.png"
    cv2.imwrite(str(raster), np.full((1000, 800, 3), 255, np.uint8))
    fp = auto_ocr._finalize_vl_page(PAGE_JSON, 1, raster)
    assert fp.image_b64
    cb = [r for r in fp.regions if r.type == doc_ir.CHECKBOX][0]
    # Ordinary OCR is immutable and does not run form interpretation implicitly.
    assert cb.markup is None

    # Geometry can still be requested as review evidence, but it is never
    # allowed to rewrite the OCR token.
    auto_ocr._apply_markup(fp, cv2.imread(str(raster)))
    assert cb.markup and cb.markup["state"] == "uncertain"
    geometry = next(
        item for item in cb.markup["observations"] if item["source"] == "geometric"
    )
    assert geometry["state"] in {"checked", "unchecked", "uncertain"}


def test_worker_protocol():
    stub = WORK / "stub.py"
    stub.write_text(
        "import json\nprint('TEXTLAB_PADDLEVL_RESULT_JSON=' + "
        "json.dumps({'pages': [{'page_number':1,'parsing_res_list':[],'layout_det_res':[],'assets':{}}]}))\n"
    )
    pages = auto_ocr.run_vl_worker([pathlib.Path("x.png")], backend_python=sys.executable, worker_path=stub)
    assert len(pages) == 1
    bad = WORK / "bad.py"
    bad.write_text("import sys; sys.exit(3)\n")
    try:
        auto_ocr.run_vl_worker([pathlib.Path("x.png")], backend_python=sys.executable, worker_path=bad)
        raise AssertionError("should raise")
    except RuntimeError:
        pass


def test_native_lane():
    _img = np.full((20, 20, 3), 128, np.uint8)
    _ok, _enc = cv2.imencode(".png", _img)
    PNG = _enc.tobytes()

    class FR:
        width = 595.0
        height = 842.0

    class FPix:
        def tobytes(self, fmt="png"):
            return PNG

    class FPage:
        rect = FR()

        def get_text(self, mode):
            if mode == "text":
                return "This is a born digital page with plenty of words in it."
            if mode == "words":
                return [("w",)] * 11
            if mode == "dict":
                return {"blocks": [
                    {"type": 0, "bbox": [72, 72, 500, 96], "lines": [{"spans": [{"text": "Hello world", "font": "Arial"}]}]},
                    {"type": 1, "bbox": [72, 120, 300, 300], "image": PNG, "ext": "png"},
                ]}
            return []

        def get_pixmap(self, dpi=150):
            return FPix()

    assert auto_ocr._page_has_text_layer(FPage()) is True
    assert auto_ocr._page_has_math(FPage()) is False
    np_page = auto_ocr._native_page(FPage(), 3)
    assert [r.type for r in np_page.regions] == ["text", "figure"]
    assert np_page.regions[0].text == "Hello world"
    assert np_page.image_b64


def test_api_compat():
    assert md.classify_from_b64("☑", None)["state"] == "checked"


def test_worker_reports_pages_as_they_finish():
    """Page progress must arrive *during* the run, not after it."""
    stub = WORK / "stub_progress.py"
    stub.write_text(
        "import json, sys, time\n"
        "total = 3\n"
        "print('TEXTLAB_PADDLEVL_PROGRESS=0/%d' % total, flush=True)\n"
        "pages = []\n"
        "for i in range(1, total + 1):\n"
        "    time.sleep(0.2)\n"
        "    pages.append({'page_number': i, 'parsing_res_list': [],\n"
        "                  'layout_det_res': [], 'assets': {}})\n"
        "    print('TEXTLAB_PADDLEVL_PROGRESS=%d/%d' % (i, total), flush=True)\n"
        "print('TEXTLAB_PADDLEVL_RESULT_JSON=' + json.dumps({'pages': pages}))\n"
    )
    seen = []
    pages = auto_ocr.run_vl_worker(
        [pathlib.Path(f"p{i}.png") for i in range(3)],
        backend_python=sys.executable, worker_path=stub,
        on_page=lambda done, total: seen.append((done, total, time.time())),
    )
    assert len(pages) == 3
    assert [(d, t) for d, t, _ in seen] == [(0, 3), (1, 3), (2, 3), (3, 3)], seen
    # Streamed, not replayed at the end: the callbacks are spread over the run.
    assert seen[-1][2] - seen[0][2] > 0.4, "progress arrived all at once"


def test_wedged_worker_is_killed_instead_of_hanging():
    """A worker that goes silent must fail loudly, not spin forever.

    This is what left the OCR page on "Running PaddleOCR-VL..." with no way out
    but reloading the browser.
    """
    stub = WORK / "stub_wedged.py"
    stub.write_text("import time\ntime.sleep(600)\n")
    started = time.time()
    try:
        auto_ocr.run_vl_worker(
            [pathlib.Path("x.png")], backend_python=sys.executable,
            worker_path=stub, stall_timeout=2.0,
        )
        raise AssertionError("a wedged worker should raise")
    except RuntimeError as exc:
        assert "no output" in str(exc)
    elapsed = time.time() - started
    assert elapsed < 30, f"took {elapsed:.1f}s to give up"


def test_a_slow_but_talking_worker_is_not_killed():
    """The guard is a stall budget, not a total one: slow pages are legitimate."""
    stub = WORK / "stub_slow.py"
    stub.write_text(
        "import json, time\n"
        "for i in range(1, 4):\n"
        "    time.sleep(0.4)\n"
        "    print('TEXTLAB_PADDLEVL_PROGRESS=%d/3' % i, flush=True)\n"
        "print('TEXTLAB_PADDLEVL_RESULT_JSON=' + json.dumps({'pages': [\n"
        "    {'page_number': 1, 'parsing_res_list': [], 'layout_det_res': [], 'assets': {}}]}))\n"
    )
    # Total runtime (~1.2s) exceeds the stall budget; no single gap does.
    pages = auto_ocr.run_vl_worker(
        [pathlib.Path("x.png")], backend_python=sys.executable,
        worker_path=stub, stall_timeout=0.9,
    )
    assert len(pages) == 1


def test_free_gpu_waits_for_the_vram_to_be_released():
    """Regression: the next document's OCR started against a resident model.

    ``keep_alive: 0`` only schedules the unload -- it returns with all ~20 GiB
    still on the card, and the PaddleOCR-VL worker (a separate process needing
    ~8.4 GiB) then died part-way through loading its own weights.
    """
    from core import vision_enrich

    calls = {"unloaded": [], "polls": 0}
    free_after = 3

    def fake_request(base_url, path, payload=None, timeout=60.0):
        if path == "/api/generate":
            calls["unloaded"].append(payload["model"])
            assert payload["keep_alive"] == 0
            return {}
        if path == "/api/ps":
            calls["polls"] += 1
            if calls["polls"] > free_after:
                return {"models": []}
            return {"models": [{"model": "vision:20b"}, {"model": "chat:8b"}]}
        return {}

    original = vision_enrich._ollama_request
    vision_enrich._ollama_request = fake_request
    try:
        evicted = vision_enrich.free_gpu()
        assert sorted(evicted) == ["chat:8b", "vision:20b"], evicted
        assert sorted(calls["unloaded"]) == ["chat:8b", "vision:20b"]
        # It kept polling until the card was actually clear.
        assert calls["polls"] > free_after, calls

        # Nothing resident: no unload requests, no waiting.
        calls.update(unloaded=[], polls=0)
        free_after = -1
        assert vision_enrich.free_gpu() == []
        assert calls["unloaded"] == []

        # A model that never unloads must not block the app forever.
        free_after = 10**9
        calls.update(unloaded=[], polls=0)
        started = time.time()
        vision_enrich.free_gpu(timeout=0.5)
        assert time.time() - started < 10
    finally:
        vision_enrich._ollama_request = original


def test_document_run_leaves_the_vision_model_warm():
    """The model must not be evicted at the end of a document.

    It expires on its own after ``keep_alive`` and the next stage that needs the
    card frees it, so evicting here would only pay a ~60 s reload for nothing.
    """
    from core import vision_enrich

    class Client(vision_enrich.OllamaVisionClient):
        def __init__(self):
            super().__init__(model="vision:20b")
            self._prepared = True
            self.unload_requests = 0

        def _request(self, path, payload=None):
            if path == "/api/generate" and (payload or {}).get("keep_alive") == 0:
                self.unload_requests += 1
            return {}

    client = Client()
    client.close()
    assert client.unload_requests == 0, "close() evicted a still-useful model"
    assert client._prepared is False

if __name__ == "__main__":
    for name, fn in sorted(globals().items()):
        if name.startswith("test_") and callable(fn):
            fn()
            print(f"OK {name}")
    print("ALL REGRESSION TESTS PASSED")
