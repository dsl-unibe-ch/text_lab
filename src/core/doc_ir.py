"""Typed intermediate representation (IR) for the automatic OCR pipeline.

A light version of the structured-document IR: ``Document -> Page -> Region``.
Everything here is plain dataclasses that are JSON-serialisable, with no heavy
runtime dependencies (only ``pandas`` for table parsing). The PaddleOCR-VL
worker and the born-digital fast lane both produce ``Page`` objects through the
adapters below; the Streamlit UI and the batch exporter consume the resulting
``Document``.

The IR is deliberately decoupled from any single engine: a region records
where it came from (``source``), how confident the layout/recognition step was
(``confidence``), and any later analysis as separate, provenance-preserving
annotations. OCR content is never rewritten by an enrichment model.
"""

from __future__ import annotations

import base64
import io
import json
import re
import zipfile
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd

# ==========================================
#        REGION TYPE TAXONOMY
# ==========================================

TEXT = "text"
TITLE = "title"
TABLE = "table"
FIGURE = "figure"
FORMULA = "formula"
FOOTNOTE = "footnote"
HEADER = "header"
FOOTER = "footer"
SEAL = "seal"
CHECKBOX = "checkbox"
LIST = "list"
REFERENCE = "reference"
OTHER = "other"

#: Region types that carry an image crop rather than (or in addition to) text.
ASSET_TYPES = {FIGURE, SEAL, CHECKBOX}

#: Maps a raw PaddleOCR-VL / PP-DocLayoutV3 ``block_label`` to an IR region type.
#: Unknown labels fall back to :data:`OTHER` while keeping their text content.
_LABEL_TO_TYPE = {
    # plain text-ish blocks
    "text": TEXT,
    "content": TEXT,
    "plain_text": TEXT,
    "number": TEXT,
    "aside_text": TEXT,
    "algorithm": TEXT,
    "formula_number": TEXT,
    # titles / captions
    "title": TITLE,
    "doc_title": TITLE,
    "paragraph_title": TITLE,
    "abstract": TITLE,
    "figure_title": TITLE,
    "chart_title": TITLE,
    "table_title": TITLE,
    "figure_caption": TITLE,
    "table_caption": TITLE,
    # tables
    "table": TABLE,
    # figures / images / charts / stamps
    "figure": FIGURE,
    "image": FIGURE,
    "chart": FIGURE,
    "header_image": FIGURE,
    "footer_image": FIGURE,
    "seal": SEAL,
    # formulas
    "formula": FORMULA,
    "interline_equation": FORMULA,
    "equation": FORMULA,
    # structure
    "header": HEADER,
    "footer": FOOTER,
    "footnote": FOOTNOTE,
    "vision_footnote": FOOTNOTE,
    "reference": REFERENCE,
    "reference_content": REFERENCE,
    "list": LIST,
    # markup
    "checkbox": CHECKBOX,
    "check_box": CHECKBOX,
}


def label_to_type(block_label: Optional[str]) -> str:
    """Map a raw engine layout label to an IR region type."""
    if not block_label:
        return OTHER
    return _LABEL_TO_TYPE.get(str(block_label).strip().lower(), OTHER)


# ==========================================
#        DATACLASSES
# ==========================================


@dataclass
class Observation:
    """One engine's observation about a visual element."""

    source: str
    value: str
    method: str = ""
    score: Optional[float] = None
    raw: Any = None

    def to_dict(self) -> Dict[str, Any]:
        return {
            "source": self.source,
            "value": self.value,
            "method": self.method,
            "score": _num(self.score) if self.score is not None else None,
            "raw": self.raw,
        }


@dataclass
class VisualDescription:
    """Generated description of a figure/image, separate from printed text."""

    description: str
    visible_text: str = ""
    source: str = ""
    model: str = ""
    warnings: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "description": self.description,
            "visible_text": self.visible_text,
            "source": self.source,
            "model": self.model,
            "warnings": list(self.warnings),
        }


@dataclass
class FormOption:
    """One labelled response option within a question or matrix row."""

    id: str
    label: str
    state: str = "unselected"  # selected|unselected|cancelled|ambiguous
    visual_mark: str = "none"  # none|x|tick|filled|scribbled|other|uncertain
    bbox: List[float] = field(default_factory=list)
    observations: List[Observation] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    associated_text: str = ""  # respondent text linked to this marked choice
    evidence_crop_b64: Optional[str] = None  # runtime/bundle asset; omitted from JSON

    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "label": self.label,
            "state": self.state,
            "visual_mark": self.visual_mark,
            "bbox": [_num(v) for v in self.bbox] if self.bbox else [],
            "observations": [o.to_dict() for o in self.observations],
            "warnings": list(self.warnings),
            "associated_text": self.associated_text,
        }


@dataclass
class FormRow:
    """A simple question row or one row of a response matrix."""

    id: str
    label: str = ""
    options: List[FormOption] = field(default_factory=list)
    status: str = "accepted"  # accepted|recovered|needs_review|failed
    warnings: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "label": self.label,
            "options": [o.to_dict() for o in self.options],
            "status": self.status,
            "warnings": list(self.warnings),
        }


@dataclass
class FormGroup:
    """Question-level semantic response annotation spanning one or more regions."""

    id: str
    bbox: List[float] = field(default_factory=list)
    question_text: str = ""
    question_type: str = "unknown"  # single|multiple|rating|matrix|unknown
    selection_rule: str = "zero_or_more"
    rows: List[FormRow] = field(default_factory=list)
    status: str = "accepted"
    warnings: List[str] = field(default_factory=list)
    parent_question_id: str = ""
    condition_text: str = ""
    provenance: Dict[str, Any] = field(default_factory=dict)
    source_crop_b64: Optional[str] = None  # runtime/bundle asset; omitted from JSON

    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "bbox": [_num(v) for v in self.bbox] if self.bbox else [],
            "question_text": self.question_text,
            "question_type": self.question_type,
            "selection_rule": self.selection_rule,
            "rows": [r.to_dict() for r in self.rows],
            "status": self.status,
            "warnings": list(self.warnings),
            "parent_question_id": self.parent_question_id,
            "condition_text": self.condition_text,
            "provenance": dict(self.provenance),
        }


@dataclass
class Region:
    """A single laid-out block on a page."""

    id: str
    type: str
    bbox: List[float]  # [x1, y1, x2, y2] in the coordinate space of Page.image_b64
    reading_order: int
    content: Dict[str, str] = field(default_factory=dict)  # {"text"|"html"|"latex"|"markdown": ...}
    confidence: Dict[str, Optional[float]] = field(default_factory=dict)  # {"layout":.., "ocr":..}
    asset: Optional[Dict[str, str]] = None  # {"b64":.., "ext":"png", "filename":..}
    source: str = "paddleocr-vl-1.6"
    warnings: List[str] = field(default_factory=list)
    markup: Optional[Dict[str, Any]] = None  # {"state":.., "method":.., "score":..} for CHECKBOX
    visual_description: Optional[VisualDescription] = None

    # -- convenience accessors -------------------------------------------------
    @property
    def text(self) -> str:
        return (
            self.content.get("text")
            or self.content.get("markdown")
            or self.content.get("html")
            or self.content.get("latex")
            or ""
        )

    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "type": self.type,
            "bbox": [_num(v) for v in self.bbox] if self.bbox else [],
            "reading_order": self.reading_order,
            "content": self.content,
            "confidence": self.confidence,
            "asset": self.asset,
            "source": self.source,
            "warnings": self.warnings,
            "markup": self.markup,
            "visual_description": (
                self.visual_description.to_dict() if self.visual_description else None
            ),
        }


@dataclass
class Page:
    page_number: int
    regions: List[Region] = field(default_factory=list)
    width: Optional[int] = None
    height: Optional[int] = None
    image_b64: Optional[str] = None  # rendered page raster (PNG) for the layout preview
    source: str = "paddleocr-vl-1.6"
    markdown: Optional[str] = None  # engine-native markdown, if any
    form_groups: List[FormGroup] = field(default_factory=list)
    #: Invisible-PDF-layer entries in full-resolution raster pixels, produced by
    #: :mod:`searchable_pdf`. Transient: valid only while the raster exists.
    text_layer: Optional[List[Dict[str, Any]]] = None
    #: ``(width, height)`` in pixels of the raster ``text_layer`` was measured
    #: on, so the PDF writer can derive the pixel->point scale from real
    #: geometry instead of assuming a DPI. Transient, like ``text_layer``.
    raster_size: Optional[Tuple[int, int]] = None
    #: Engine that supplied the word geometry, kept for the citable summary
    #: after ``text_layer`` itself is dropped.
    text_layer_engine: str = ""

    def ordered_regions(self) -> List[Region]:
        return sorted(self.regions, key=lambda r: (r.reading_order, r.id))

    def to_dict(self) -> Dict[str, Any]:
        return {
            "page_number": self.page_number,
            "width": self.width,
            "height": self.height,
            "source": self.source,
            "regions": [r.to_dict() for r in self.regions],
            "form_groups": [g.to_dict() for g in self.form_groups],
            # image_b64 / markdown / text_layer deliberately excluded from the
            # canonical JSON to keep it compact; assets and the searchable PDF
            # are exported separately.
        }


@dataclass
class Document:
    pages: List[Page] = field(default_factory=list)
    source_name: str = ""
    #: Original pages plus an invisible text layer, when the export was
    #: requested. Bytes, so deliberately not part of :meth:`to_dict`.
    searchable_pdf: Optional[bytes] = None
    #: Tools that cannot be inferred from the regions themselves, e.g.
    #: ``{"text_layer": "Tesseract 4.1.1 (deu)"}``. Merged into the citable
    #: summary by :func:`model_provenance`.
    extra_tools: Dict[str, str] = field(default_factory=dict)

    def all_regions(self):
        for page in self.pages:
            for region in page.ordered_regions():
                yield page, region

    def to_dict(self) -> Dict[str, Any]:
        return {
            "source_name": self.source_name,
            "n_pages": len(self.pages),
            "models": model_provenance(self),
            "pages": [p.to_dict() for p in self.pages],
        }


def _num(value):
    try:
        f = float(value)
        return int(f) if f.is_integer() else round(f, 2)
    except (TypeError, ValueError):
        return value


# ==========================================
#        PADDLEOCR-VL ADAPTER
# ==========================================


def _as_bbox(raw) -> List[float]:
    """Normalise a bbox/coordinate into [x1, y1, x2, y2]."""
    if raw is None:
        return []
    try:
        flat = []
        for v in raw:
            if isinstance(v, (list, tuple)):
                flat.extend(v)
            else:
                flat.append(v)
        nums = [float(x) for x in flat]
    except (TypeError, ValueError):
        return []
    if len(nums) == 4:
        x1, y1, x2, y2 = nums
        return [min(x1, x2), min(y1, y2), max(x1, x2), max(y1, y2)]
    if len(nums) >= 8:
        xs = nums[0::2]
        ys = nums[1::2]
        return [min(xs), min(ys), max(xs), max(ys)]
    return nums


def _iou(a: List[float], b: List[float]) -> float:
    if len(a) < 4 or len(b) < 4:
        return 0.0
    ax1, ay1, ax2, ay2 = a[:4]
    bx1, by1, bx2, by2 = b[:4]
    ix1, iy1 = max(ax1, bx1), max(ay1, by1)
    ix2, iy2 = min(ax2, bx2), min(ay2, by2)
    iw, ih = max(0.0, ix2 - ix1), max(0.0, iy2 - iy1)
    inter = iw * ih
    if inter <= 0:
        return 0.0
    area_a = max(0.0, ax2 - ax1) * max(0.0, ay2 - ay1)
    area_b = max(0.0, bx2 - bx1) * max(0.0, by2 - by1)
    union = area_a + area_b - inter
    return inter / union if union > 0 else 0.0


def _match_layout_score(bbox: List[float], layout_boxes: List[dict]) -> Optional[float]:
    """Best-overlap layout confidence for a parsing block."""
    best_score = None
    best_iou = 0.0
    for box in layout_boxes or []:
        coord = _as_bbox(box.get("coordinate") or box.get("bbox"))
        iou = _iou(bbox, coord)
        if iou > best_iou:
            best_iou = iou
            score = box.get("score")
            best_score = float(score) if score is not None else None
    return best_score if best_iou >= 0.3 else None


def from_paddle_vl(page_json: Dict[str, Any]) -> Page:
    """Build a :class:`Page` from one page of PaddleOCR-VL worker output.

    Expected (best-effort) ``page_json`` keys::

        page_number, width, height, image_b64, markdown,
        parsing_res_list: [{block_label, block_content, block_bbox, block_order}],
        layout_det_res:   [{label, score, coordinate}],
        assets:           {"<index>": {"b64":.., "ext":"png"}}
    """
    page = Page(
        page_number=int(page_json.get("page_number") or page_json.get("page") or 1),
        width=page_json.get("width"),
        height=page_json.get("height"),
        image_b64=page_json.get("image_b64"),
        source="paddleocr-vl-1.6",
        markdown=page_json.get("markdown"),
    )

    layout_boxes = page_json.get("layout_det_res") or []
    assets = page_json.get("assets") or {}
    blocks = page_json.get("parsing_res_list") or []

    for idx, block in enumerate(blocks):
        label = block.get("block_label") or block.get("label")
        rtype = label_to_type(label)
        bbox = _as_bbox(block.get("block_bbox") or block.get("bbox"))
        raw_content = block.get("block_content")
        if raw_content is None:
            raw_content = block.get("content", "")
        content_str = "" if raw_content is None else str(raw_content)

        content: Dict[str, str] = {}
        if rtype == TABLE:
            content["html"] = content_str
        elif rtype == FORMULA:
            content["latex"] = content_str
        else:
            content["text"] = content_str
        if content_str and rtype not in (TABLE, FORMULA):
            content["markdown"] = content_str

        order = block.get("block_order")
        reading_order = int(order) if isinstance(order, (int, float)) else idx

        asset = None
        for key in (str(idx), idx):
            if key in assets:
                asset = assets[key]
                break
        if asset is None and block.get("block_image_b64"):
            asset = {"b64": block["block_image_b64"], "ext": "png"}

        region = Region(
            id=f"p{page.page_number}_r{idx}",
            type=rtype,
            bbox=bbox,
            reading_order=reading_order,
            content=content,
            confidence={
                "layout": _match_layout_score(bbox, layout_boxes),
                "ocr": _num(block["block_score"]) if block.get("block_score") is not None else None,
            },
            asset=asset,
            source="paddleocr-vl-1.6",
        )
        if rtype == OTHER and label:
            region.warnings.append(f"unmapped layout label: {label}")
        page.regions.append(region)

    return page


# ==========================================
#        TABLE PARSING
# ==========================================


def extract_html_table(html_content: str) -> Optional[pd.DataFrame]:
    """Parse the first HTML ``<table>`` in *html_content* into a DataFrame.

    Generalised, per-region version of the whole-text regex that used to live
    in ``ocr_engine.py``. Returns ``None`` if there is no parseable table.
    """
    if not html_content:
        return None
    if "<table" not in html_content:
        return None
    try:
        dfs = pd.read_html(io.StringIO(html_content))
    except Exception:
        return None
    return dfs[0] if dfs else None


def tables_to_dataframes(document: Document) -> List[Dict[str, Any]]:
    """Return one entry per parseable table region across the whole document."""
    tables = []
    for page, region in document.all_regions():
        if region.type != TABLE:
            continue
        html = region.content.get("html", "")
        df = extract_html_table(html)
        if df is None:
            continue
        tables.append(
            {
                "page": page.page_number,
                "region_id": region.id,
                "dataframe": df,
                "html": html,
            }
        )
    return tables


# ==========================================
#        MARKDOWN / ASSET EXPORT
# ==========================================


def _safe_filename(region: Region) -> str:
    ext = (region.asset or {}).get("ext", "png")
    slug = re.sub(r"[^A-Za-z0-9_.-]", "_", region.id)
    return f"{slug}.{ext}"


def _safe_slug(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]", "_", str(value or "item"))


def _region_to_markdown(region: Region, asset_dir: str, embed_assets: bool) -> str:
    rtype = region.type
    if rtype == TITLE:
        text = region.text.strip()
        return f"## {text}" if text else ""
    if rtype == TABLE:
        html = region.content.get("html", "").strip()
        return html
    if rtype == FORMULA:
        latex = region.content.get("latex", "").strip()
        return f"$$\n{latex}\n$$" if latex else ""
    if rtype == CHECKBOX:
        state = (region.markup or {}).get("state", "uncertain")
        marker = {"checked": "[x]", "unchecked": "[ ]", "uncertain": "[?]"}.get(state, "[?]")
        label = region.text.strip()
        return f"- {marker} {label}".rstrip()
    if rtype in ASSET_TYPES and region.asset and embed_assets:
        fname = _safe_filename(region)
        caption = region.text.strip()
        img = f"![{caption or region.id}]({asset_dir}/{fname})"
        return f"{img}\n\n{caption}" if caption else img
    # text-ish fallback
    return region.text.strip()


def to_markdown(document: Document, asset_dir: str = "assets", embed_assets: bool = True) -> str:
    """Deterministic markdown rendering of the whole document (reading order)."""
    chunks: List[str] = []
    for page in document.pages:
        if len(document.pages) > 1:
            chunks.append(f"<!-- page {page.page_number} -->")
        for region in page.ordered_regions():
            md = _region_to_markdown(region, asset_dir, embed_assets)
            if md.strip():
                chunks.append(md)
    return "\n\n".join(chunks).strip() + "\n"


def _region_to_text(region: Region) -> str:
    """Plain-text rendering: no markup, but structure kept legible."""
    rtype = region.type
    if rtype == TABLE:
        df = extract_html_table(region.content.get("html", ""))
        if df is None:
            return region.text.strip()
        return df.to_string(index=False)
    if rtype == FORMULA:
        return region.content.get("latex", "").strip()
    if rtype == CHECKBOX:
        state = (region.markup or {}).get("state", "uncertain")
        marker = {"checked": "[x]", "unchecked": "[ ]", "uncertain": "[?]"}.get(state, "[?]")
        return f"{marker} {region.text.strip()}".rstrip()
    if rtype in ASSET_TYPES and region.asset:
        caption = region.text.strip()
        return f"[{rtype}: {caption}]" if caption else f"[{rtype}]"
    return region.text.strip()


def to_text(document: Document, page_separator: bool = True) -> str:
    """Plain-text rendering of the whole document in reading order."""
    chunks: List[str] = []
    for page in document.pages:
        if page_separator and len(document.pages) > 1:
            chunks.append(f"--- page {page.page_number} ---")
        for region in page.ordered_regions():
            text = _region_to_text(region)
            if text.strip():
                chunks.append(text)
    return "\n\n".join(chunks).strip() + "\n"


def build_docx(document: Document, doc_stem: str = "document") -> Optional[bytes]:
    """An editable .docx with headings, paragraphs, real tables and figures.

    Returns ``None`` when ``python-docx`` is unavailable, so the caller can hide
    the download instead of failing the whole export.
    """
    try:
        import docx
        from docx.shared import Inches, Pt
    except ImportError:
        return None

    doc = docx.Document()
    for page_index, page in enumerate(document.pages):
        if page_index:
            doc.add_page_break()
        for region in page.ordered_regions():
            rtype = region.type
            if rtype == TITLE:
                text = region.text.strip()
                if text:
                    doc.add_heading(text, level=2)
            elif rtype == TABLE:
                df = extract_html_table(region.content.get("html", ""))
                if df is None:
                    if region.text.strip():
                        doc.add_paragraph(region.text.strip())
                    continue
                table = doc.add_table(rows=1, cols=max(len(df.columns), 1))
                table.style = "Table Grid"
                for cell, column in zip(table.rows[0].cells, df.columns):
                    run = cell.paragraphs[0].add_run(str(column))
                    run.bold = True
                for record in df.itertuples(index=False):
                    for cell, value in zip(table.add_row().cells, record):
                        cell.text = "" if pd.isna(value) else str(value)
            elif rtype == FORMULA:
                latex = region.content.get("latex", "").strip()
                if latex:
                    run = doc.add_paragraph().add_run(latex)
                    run.font.name = "Consolas"
                    run.font.size = Pt(10)
            elif rtype == CHECKBOX:
                state = (region.markup or {}).get("state", "uncertain")
                marker = {"checked": "☒", "unchecked": "☐", "uncertain": "☐?"}.get(state, "☐?")
                doc.add_paragraph(f"{marker} {region.text.strip()}".rstrip())
            elif rtype in ASSET_TYPES and region.asset and region.asset.get("b64"):
                try:
                    doc.add_picture(
                        io.BytesIO(base64.b64decode(region.asset["b64"])), width=Inches(5.5)
                    )
                except Exception:
                    doc.add_paragraph(f"[{rtype} could not be embedded]")
                caption = region.text.strip()
                if caption:
                    # Italics live on the run, not the paragraph: assigning
                    # ``Paragraph.italic`` is accepted silently and does nothing.
                    doc.add_paragraph().add_run(caption).italic = True
            elif region.text.strip():
                doc.add_paragraph(region.text.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def collect_assets(document: Document) -> Dict[str, bytes]:
    """Return ``{filename: png_bytes}`` for every region that carries a crop."""
    out: Dict[str, bytes] = {}
    for _, region in document.all_regions():
        if region.type not in ASSET_TYPES or not region.asset:
            continue
        b64 = region.asset.get("b64")
        if not b64:
            continue
        try:
            out[_safe_filename(region)] = base64.b64decode(b64)
        except Exception:
            continue
    for page in document.pages:
        for group in page.form_groups:
            if group.source_crop_b64:
                try:
                    out[f"form_{_safe_slug(group.id)}.png"] = base64.b64decode(
                        group.source_crop_b64
                    )
                except Exception:
                    pass
            for row in group.rows:
                for option in row.options:
                    if not option.evidence_crop_b64:
                        continue
                    try:
                        name = (
                            f"form_{_safe_slug(group.id)}_"
                            f"{_safe_slug(row.id)}_{_safe_slug(option.id)}.png"
                        )
                        out[name] = base64.b64decode(option.evidence_crop_b64)
                    except Exception:
                        pass
    return out


def form_responses_to_dataframe(document: Document) -> Optional[pd.DataFrame]:
    """Flatten semantic form responses into one row per question/matrix row."""
    def _answer(option: FormOption) -> str:
        if not option.associated_text:
            return option.label
        return f"{option.label.rstrip().rstrip(':')}: {option.associated_text}"

    records: List[Dict[str, Any]] = []
    for page in document.pages:
        for group in page.form_groups:
            for row in group.rows:
                selected = [_answer(o) for o in row.options if o.state == "selected"]
                cancelled = [_answer(o) for o in row.options if o.state == "cancelled"]
                ambiguous = [_answer(o) for o in row.options if o.state == "ambiguous"]
                records.append(
                    {
                        "page": page.page_number,
                        "group_id": group.id,
                        "parent_group_id": group.parent_question_id,
                        "condition": group.condition_text,
                        "question": group.question_text,
                        "row_id": row.id,
                        "row": row.label,
                        "selected": " | ".join(selected),
                        "cancelled": " | ".join(cancelled),
                        "ambiguous": " | ".join(ambiguous),
                        "status": (
                            "needs_review"
                            if "needs_review" in (group.status, row.status)
                            else row.status
                        ),
                    }
                )
    return pd.DataFrame(records) if records else None


def build_form_responses_csv(document: Document) -> Optional[bytes]:
    """Semantic form-response CSV, or ``None`` if no groups were extracted."""
    df = form_responses_to_dataframe(document)
    if df is None:
        return None
    return df.to_csv(index=False).encode("utf-8")


#: Human-readable, citable names for the ``Page.source`` tags the lanes record.
_SOURCE_CITATIONS = {
    "paddleocr-vl-1.6": "PaddleOCR-VL 1.6 (PaddleOCR 3.7 / PaddleX 3.7)",
    "native": "PyMuPDF text extraction (no recognition model)",
}


def model_provenance(document: Document) -> Dict[str, Any]:
    """Which models produced this result, for citation in a publication.

    Read off what the pipeline already recorded -- the per-page lane tag and the
    model named on each generated description -- so it cannot drift from what
    actually ran. A researcher quoting the OCR needs the recogniser; one quoting
    a generated figure description needs that model too, and needs to be able to
    tell the two apart.
    """
    recognition: List[str] = []
    for page in document.pages:
        name = _SOURCE_CITATIONS.get(page.source, page.source)
        if name and name not in recognition:
            recognition.append(name)

    descriptions: List[str] = []
    for _, region in document.all_regions():
        described = region.visual_description
        if described is None or not described.model:
            continue
        name = f"{described.model} ({described.source})" if described.source else described.model
        if name not in descriptions:
            descriptions.append(name)

    models: Dict[str, Any] = {"text_recognition": recognition}
    if descriptions:
        models["figure_descriptions"] = descriptions
    for key, value in (document.extra_tools or {}).items():
        if value:
            models[key] = value
    return models


def model_provenance_text(document: Document) -> str:
    """The provenance summary as lines suitable for a README or a caption."""
    labels = {
        "text_recognition": "Text recognition",
        "figure_descriptions": "Figure descriptions",
        "text_layer": "Searchable-PDF word geometry",
    }
    lines = []
    for key, value in model_provenance(document).items():
        if not value:
            continue
        joined = ", ".join(value) if isinstance(value, list) else str(value)
        lines.append(f"{labels.get(key, key.replace('_', ' ').capitalize())}: {joined}")
    return "\n".join(lines)


def to_json(document: Document, indent: int = 2) -> str:
    return json.dumps(document.to_dict(), ensure_ascii=False, indent=indent)


# ==========================================
#        ZIP BUNDLING
# ==========================================


def build_markdown_zip(document: Document, doc_stem: str = "document") -> bytes:
    """A ZIP with ``document.md`` plus an ``assets/`` folder of crops."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{doc_stem}.md", to_markdown(document))
        for fname, data in collect_assets(document).items():
            zf.writestr(f"assets/{fname}", data)
    return buf.getvalue()


def build_tables_csv_zip(document: Document) -> Optional[bytes]:
    """A ZIP of one CSV per parseable table region, or ``None`` if no tables."""
    tables = tables_to_dataframes(document)
    if not tables:
        return None
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for entry in tables:
            csv_bytes = entry["dataframe"].to_csv(index=False).encode("utf-8")
            zf.writestr(f"tables/table_{entry['region_id']}.csv", csv_bytes)
    return buf.getvalue()


def build_full_bundle(document: Document, doc_stem: str = "document") -> bytes:
    """Everything: markdown, plain text, .docx, canonical JSON, assets, table CSVs."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{doc_stem}.md", to_markdown(document))
        zf.writestr(f"{doc_stem}.txt", to_text(document))
        zf.writestr(f"{doc_stem}.json", to_json(document))
        docx_bytes = build_docx(document, doc_stem)
        if docx_bytes:
            zf.writestr(f"{doc_stem}.docx", docx_bytes)
        if document.searchable_pdf:
            zf.writestr(f"{doc_stem}_searchable.pdf", document.searchable_pdf)
        for fname, data in collect_assets(document).items():
            zf.writestr(f"assets/{fname}", data)
        for entry in tables_to_dataframes(document):
            csv_bytes = entry["dataframe"].to_csv(index=False).encode("utf-8")
            zf.writestr(f"tables/table_{entry['region_id']}.csv", csv_bytes)
        form_csv = build_form_responses_csv(document)
        if form_csv:
            zf.writestr("responses/form_responses.csv", form_csv)
        provenance = model_provenance_text(document)
        if provenance:
            zf.writestr("models_used.txt", provenance + "\n")
    return buf.getvalue()


def write_document_outputs(document: Document, out_dir, stem: str = "document") -> List[str]:
    """Write every export format for one document into *out_dir*.

    Shared with the batch runner so a format added to the single-document
    downloads cannot silently go missing from a batch result: both go through
    the same list. Returns the relative names written.
    """
    import os

    out_dir = str(out_dir)
    os.makedirs(out_dir, exist_ok=True)
    written: List[str] = []

    def _write(name: str, data, binary: bool = False):
        path = os.path.join(out_dir, name)
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "wb" if binary else "w", **({} if binary else {"encoding": "utf-8"})) as fh:
            fh.write(data)
        written.append(name)

    _write(f"{stem}.md", to_markdown(document))
    _write(f"{stem}.txt", to_text(document))
    _write(f"{stem}.json", to_json(document))
    docx_bytes = build_docx(document, stem)
    if docx_bytes:
        _write(f"{stem}.docx", docx_bytes, binary=True)
    if document.searchable_pdf:
        _write(f"{stem}_searchable.pdf", document.searchable_pdf, binary=True)
    for fname, data in collect_assets(document).items():
        _write(os.path.join("assets", fname), data, binary=True)
    for entry in tables_to_dataframes(document):
        _write(
            os.path.join("tables", f"table_{entry['region_id']}.csv"),
            entry["dataframe"].to_csv(index=False),
        )
    form_csv = build_form_responses_csv(document)
    if form_csv:
        _write("form_responses.csv", form_csv, binary=True)
    provenance = model_provenance_text(document)
    if provenance:
        _write("models_used.txt", provenance + "\n")
    return written
